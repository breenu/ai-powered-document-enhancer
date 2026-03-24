"""Document formatting module with keyword-based type detection and templates.

Detects document type (essay, report, letter, notes, research paper) via
keyword scoring, then applies python-docx formatting templates.
"""

import json
import logging
import os
import sys
from dataclasses import dataclass
from pathlib import Path
from typing import Dict, List, Optional

from docx import Document as DocxDocument
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

logger = logging.getLogger(__name__)

_ALIGNMENT_MAP = {
    "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
    "left": WD_ALIGN_PARAGRAPH.LEFT,
    "center": WD_ALIGN_PARAGRAPH.CENTER,
    "right": WD_ALIGN_PARAGRAPH.RIGHT,
}


TYPE_KEYWORDS: Dict[str, List[str]] = {
    "essay": [
        "essay", "thesis statement", "in conclusion", "argue",
        "argument", "furthermore", "in summary",
    ],
    "report": [
        "report", "findings", "methodology", "results", "analysis",
        "executive summary", "recommendation",
    ],
    "letter": [
        "dear", "sincerely", "regards", "to whom it may concern",
        "yours faithfully", "enclosed",
    ],
    "notes": [
        "note", "bullet", "key point", "reminder", "todo",
        "action item", "meeting notes",
    ],
    "research_paper": [
        "abstract", "introduction", "literature review", "methodology",
        "hypothesis", "references", "citation", "doi",
    ],
}


@dataclass
class FormatTemplate:
    name: str
    font_name: str = "Times New Roman"
    font_size: int = 12
    line_spacing: float = 1.5
    heading_font_size: int = 16
    margin_inches: float = 1.0
    alignment: int = WD_ALIGN_PARAGRAPH.JUSTIFY
    first_line_indent: float = 0.5
    header_text: str = ""
    footer_text: str = ""


_FALLBACK_TEMPLATES: Dict[str, FormatTemplate] = {
    "essay": FormatTemplate(
        name="Essay",
        font_name="Times New Roman",
        font_size=12,
        line_spacing=2.0,
        first_line_indent=0.5,
        alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
    ),
    "report": FormatTemplate(
        name="Report",
        font_name="Calibri",
        font_size=11,
        line_spacing=1.5,
        heading_font_size=18,
        first_line_indent=0.0,
        alignment=WD_ALIGN_PARAGRAPH.LEFT,
    ),
    "letter": FormatTemplate(
        name="Letter",
        font_name="Arial",
        font_size=12,
        line_spacing=1.15,
        first_line_indent=0.0,
        alignment=WD_ALIGN_PARAGRAPH.LEFT,
    ),
    "notes": FormatTemplate(
        name="Notes",
        font_name="Calibri",
        font_size=11,
        line_spacing=1.15,
        first_line_indent=0.0,
        alignment=WD_ALIGN_PARAGRAPH.LEFT,
    ),
    "research_paper": FormatTemplate(
        name="Research Paper",
        font_name="Times New Roman",
        font_size=12,
        line_spacing=2.0,
        heading_font_size=14,
        first_line_indent=0.5,
        alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
    ),
}


def _template_dir() -> Path:
    """Resolve the resources/templates/ directory for dev and PyInstaller."""
    if getattr(sys, "_MEIPASS", None):
        return Path(sys._MEIPASS) / "resources" / "templates"
    return Path(__file__).resolve().parent.parent.parent / "resources" / "templates"


def _parse_template_json(data: dict) -> FormatTemplate:
    alignment = _ALIGNMENT_MAP.get(
        str(data.get("alignment", "justify")).lower(),
        WD_ALIGN_PARAGRAPH.JUSTIFY,
    )
    return FormatTemplate(
        name=data.get("name", ""),
        font_name=data.get("font_name", "Times New Roman"),
        font_size=int(data.get("font_size", 12)),
        line_spacing=float(data.get("line_spacing", 1.5)),
        heading_font_size=int(data.get("heading_font_size", 16)),
        margin_inches=float(data.get("margin_inches", 1.0)),
        alignment=alignment,
        first_line_indent=float(data.get("first_line_indent", 0.5)),
        header_text=data.get("header_text", ""),
        footer_text=data.get("footer_text", ""),
    )


def load_templates() -> Dict[str, FormatTemplate]:
    """Load templates from resources/templates/*.json, falling back to built-in defaults."""
    templates: Dict[str, FormatTemplate] = {}
    tpl_dir = _template_dir()

    if tpl_dir.is_dir():
        for json_file in sorted(tpl_dir.glob("*.json")):
            try:
                data = json.loads(json_file.read_text(encoding="utf-8"))
                key = json_file.stem
                templates[key] = _parse_template_json(data)
                logger.debug("Loaded template '%s' from %s", key, json_file)
            except (json.JSONDecodeError, KeyError, ValueError) as exc:
                logger.warning("Skipping invalid template %s: %s", json_file.name, exc)

    if not templates:
        logger.info("No external templates found; using built-in defaults")
        return dict(_FALLBACK_TEMPLATES)

    for key, fallback in _FALLBACK_TEMPLATES.items():
        templates.setdefault(key, fallback)

    return templates


DEFAULT_TEMPLATES: Dict[str, FormatTemplate] = load_templates()


class DocumentFormatter:
    """Detects document type and applies formatting templates via python-docx."""

    def __init__(self, templates: Optional[Dict[str, FormatTemplate]] = None):
        self.templates = templates or DEFAULT_TEMPLATES

    def detect_type(self, text: str) -> str:
        text_lower = text.lower()
        scores: Dict[str, int] = {}
        for doc_type, keywords in TYPE_KEYWORDS.items():
            scores[doc_type] = sum(1 for kw in keywords if kw in text_lower)

        if max(scores.values(), default=0) == 0:
            return "essay"

        return max(scores, key=scores.get)

    def apply_template(self, text: str, doc_type: str = None,
                       output_path: str = None) -> DocxDocument:
        if doc_type is None:
            doc_type = self.detect_type(text)

        template = self.templates.get(doc_type, self.templates["essay"])
        doc = DocxDocument()

        for section in doc.sections:
            section.top_margin = Inches(template.margin_inches)
            section.bottom_margin = Inches(template.margin_inches)
            section.left_margin = Inches(template.margin_inches)
            section.right_margin = Inches(template.margin_inches)

        paragraphs = text.split("\n")

        for para_text in paragraphs:
            stripped = para_text.strip()
            if not stripped:
                doc.add_paragraph("")
                continue

            is_heading = (
                stripped.isupper()
                or stripped.endswith(":")
                or (len(stripped.split()) <= 6 and not stripped.endswith("."))
            )

            if is_heading and len(stripped) < 100:
                heading = doc.add_heading(stripped, level=1)
                for run in heading.runs:
                    run.font.size = Pt(template.heading_font_size)
                    run.font.name = template.font_name
            else:
                paragraph = doc.add_paragraph()
                paragraph.alignment = template.alignment
                pf = paragraph.paragraph_format
                pf.line_spacing = template.line_spacing
                if template.first_line_indent > 0:
                    pf.first_line_indent = Inches(template.first_line_indent)
                run = paragraph.add_run(stripped)
                run.font.name = template.font_name
                run.font.size = Pt(template.font_size)

        if output_path:
            os.makedirs(os.path.dirname(output_path) or ".", exist_ok=True)
            doc.save(output_path)
            logger.info("Formatted document saved to %s", output_path)

        return doc

    def get_template_names(self) -> List[str]:
        return list(self.templates.keys())
